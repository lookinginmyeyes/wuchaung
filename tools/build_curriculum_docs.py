from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "frontend" / "assets" / "docs"
LECTURE_PATH = DOC_DIR / "ai-vision-falling-ball-lecture.docx"
QUIZ_PATH = DOC_DIR / "ai-vision-falling-ball-quiz.docx"


LECTURE_SECTIONS = [
    (
        "一、实验目标与平台流程",
        [
            "本实验使用落球法测量液体粘滞系数。平台的作用不是替代实验判断，而是帮助学生把小球下落轨迹、终端速度、修正因子和误差来源显示出来。",
            "完整流程为：阅读讲义，完成准入题目；进入 AI 实验模块并填写实验参数；连接实时画面；标注量筒边缘并完成标定；标定后关灯、打开背光光源；释放小球并进行 AI 视觉测量，同时完成必要的人工测量；把人工测量值输入平台；最后生成实验报告并复核结果。",
            "实验前必须明确并记录液体密度、小球半径、小球密度、量筒内径、液体深度、温度等参数。平台中的参考液体只提供参考粘度、液体温度和液体密度，其余几何量仍需由实验者自己测量和填写。",
        ],
    ),
    (
        "二、落球法的物理依据",
        [
            "半径为 r 的光滑小球在静止液体中下落时，受到重力、浮力和液体粘滞阻力。刚释放时速度较小，阻力较小，小球加速下落；速度增大后阻力增大。当重力等于浮力与阻力之和时，小球加速度趋近于零，进入终端匀速阶段。",
            "在低速、低雷诺数、无明显涡流的条件下，斯托克斯定律给出小球受到的粘滞阻力：F = 6πηrv。",
            "小球达到终端速度 vt 后，可由受力平衡得到理想 Stokes 公式：η₀ = 2r²g(ρ球 − ρ液) / (9vt)。其中 ρ球 为小球密度，ρ液 为液体密度，g 取 9.80665 m/s²。",
            "若还未进入匀速段就取速度，vt 通常偏小，计算得到的 η 会偏大。",
        ],
    ),
    (
        "三、适用条件与修正",
        [
            "小球应刚性、光滑、近似球形，释放时尽量无初速度、无旋转、无明显偏斜。",
            "液体应近似为牛顿流体，且实验温度稳定；粘度对温度敏感，必须记录温度。",
            "小球运动处于低雷诺数状态，教学实验中通常要求 Re < 1，更严格时可参考 Re < 0.1。",
            "小球应在量筒中心轴线附近下落，避免贴壁；球筒径比 2r/D 越小，壁面影响越弱。",
            "取值区间必须位于终端匀速段，不能把液面附近加速段或筒底附近受扰动区间用于主计算。",
            "真实量筒不是无限宽广液体，筒壁和液柱深度会改变小球周围液体流动。平台保留壁效应和雷诺数修正，并同时显示理想结果与修正结果，便于比较。",
            "修正关系包括：Re = ρ液·vt·(2r)/η，K壁 = (1 + 2.4r/R)(1 + 3.3r/H)，KRe = 1 + 3Re/16 − 19Re²/1080。式中 R 为量筒内半径，H 为待测液体深度。",
            "修正项不能消除所有误差，它只是把已知系统偏差纳入模型；若 Re、壁效应、释放偏斜或追踪置信度明显不合格，应优先重测。",
        ],
    ),
    (
        "四、标定与视觉测量",
        [
            "AI 测量依赖摄像头画面中的像素坐标，因此必须先完成标定。平台要求先标注量筒左右边缘，系统生成量筒中心虚线；开始标定后，应沿这条中心线从上到下点击标定棒刻度点。",
            "标定棒的半径、长度和刻度间距由实验者手动输入，平台根据这些数据自动显示需要点击的点数。若点错，可点击该点删除后重新标定。",
            "标定完成后，不能移动相机、量筒和释放装置；否则像素坐标与真实坐标的对应关系会失效。标定棒应取出后再开始落球测量，避免遮挡小球。",
            "本实验推荐背光成像：标定完成后关闭环境灯，打开背光光源，使小球在画面中形成稳定、清晰的暗色轮廓。背光应均匀，摄像头曝光不宜过高；过曝会让小球边缘发白，阴影、气泡和划痕更容易被误识别。",
        ],
    ),
    (
        "五、AI 测量与人工测量",
        [
            "AI 测量时，平台逐帧识别小球中心，生成 s-t 图和 v-t 图，并从速度较稳定、异常点较少的区间拟合终端速度 vt。",
            "结果区会显示终端速度、理想粘滞系数、修正粘滞系数、Re、壁效应修正、追踪置信度和匀速段长度。",
            "人工测量不是可有可无的对照。实验者应根据标定长度或量筒刻度，记录小球通过选定区间的距离和时间，估算人工终端速度与人工粘滞系数，并把两个值输入平台。",
            "平台会根据 AI 结果与人工结果的差异生成评分和报告，帮助判断误差主要来自释放、标定、计时、选段还是视觉追踪。",
        ],
    ),
    (
        "六、推荐操作步骤",
        [
            "1. 阅读讲义并完成准入题目，通过后进入实验大厅。",
            "2. 在 AI 实验模块填写实验参数：液体密度、小球半径、量筒内径、液体深度、小球密度和温度等。",
            "3. 连接实时画面，固定手机或摄像头，使量筒竖直、中心线清楚，小球下落区域位于画面中央。",
            "4. 标注量筒左右边缘，确认中心虚线；把标定棒放在量筒中心轴线，输入标定棒半径、长度和刻度间距。",
            "5. 点击开始标定，在中心虚线上按真实刻度顺序点击标定点；若点错，可点击该点删除后重新标定。",
            "6. 标定完成后取出标定棒，关闭环境灯，打开背光光源，调整曝光使小球轮廓清晰且不过曝。",
            "7. 点击开始实时追踪，静止释放小球；若小球偏离中心线、贴壁、气泡被误识别或轨迹断裂，应停止并重测。",
            "8. 完成 AI 追踪后，同时完成或录入人工测量值：人工终端速度和人工粘滞系数。",
            "9. 查看平台生成的曲线、修正结果、诊断建议和网页报告，必要时载入历史记录进行复盘比较。",
        ],
    ),
    (
        "七、结果复核与误差判断",
        [
            "先看 s-t 图是否整体单调、平滑，v-t 图是否出现持续稳定的终端速度平台。",
            "检查 Re 是否满足低雷诺数要求，若 Re 接近或超过 1，应考虑换更小小球、更高粘度液体或重新设计实验条件。",
            "检查 2r/D 和 K壁。小球越大、量筒越窄，壁效应越明显，修正结果与理想结果差异也可能变大。",
            "检查追踪置信度和匀速段长度。若误识别气泡、划痕或阴影，应重新框选检测区域、调整背光和曝光，必要时重测。",
            "比较人工测量值与 AI 测量值。若差异较大，应从人工计时、标定点、释放中心性、背光条件和算法追踪质量逐项排查。",
        ],
    ),
]


QUIZ_ITEMS = [
    ("单项选择题", "本平台完整实验流程的合理顺序是（ ）", "B", "A. 先测量，再读讲义，最后标定\nB. 阅读讲义与答题准入后，填写参数、连接画面、标定、背光测量、录入人工值并生成报告\nC. 只要连接摄像头即可直接生成报告\nD. 先关闭背光，再进行标定和读题"),
    ("单项选择题", "落球法计算液体粘滞系数最关键的速度量是（ ）", "A", "A. 小球进入稳定匀速阶段后的终端速度 vt\nB. 刚释放瞬间的速度\nC. 液面附近的最大加速度\nD. 任意一帧画面中的瞬时速度"),
    ("单项选择题", "斯托克斯定律 F=6πηrv 主要适用于（ ）", "B", "A. 小球高速运动\nB. 小球低速运动且液体无涡流\nC. 液体产生强烈涡流\nD. 大体积球体运动"),
    ("判断题", "若小球尚未进入匀速段就取速度，通常会使终端速度偏小，并可能使计算出的粘滞系数偏大。", "正确", ""),
    ("判断题", "教学实验中通常要求雷诺数 Re < 1；若希望更接近蠕动流条件，可采用 Re < 0.1 作为更保守判据。", "正确", ""),
    ("单项选择题", "下列哪种做法最有助于减小壁效应（ ）", "A", "A. 选用更大内径的量筒或更小的小球，使 2r/D 变小\nB. 让小球贴近筒壁下落\nC. 提高环境光亮度使画面更白\nD. 不记录量筒内径"),
    ("单项选择题", "开始标定前，平台先标注量筒左右边缘并生成中心虚线，主要目的是（ ）", "B", "A. 让画面更好看\nB. 确定量筒中心轴线，使标定点和落球路径尽量沿中心线\nC. 自动改变液体密度\nD. 替代小球半径测量"),
    ("判断题", "标定完成后可以随意移动摄像头或量筒，因为平台会自动保持原来的像素-长度关系。", "错误", ""),
    ("判断题", "背光测量时，摄像头曝光越高越好，过曝不会影响小球边缘识别。", "错误", ""),
    ("单项选择题", "平台判定终端速度时，最应关注的图像和曲线特征是（ ）", "B", "A. 小球颜色是否最深\nB. v-t 图中是否出现持续稳定的平台段\nC. 画面是否越亮越好\nD. 第一帧速度是否最大"),
    ("单项选择题", "若 AI 测量结果和人工测量结果差异较大，最合理的处理是（ ）", "D", "A. 直接删除人工测量值\nB. 只相信 AI 结果，不再复核\nC. 只相信人工结果，不看曲线\nD. 检查人工计时、标定点、释放中心性、背光曝光和追踪质量"),
    ("判断题", "如果画面中气泡或量筒划痕被反复误识别，应调整检测区域、背光和曝光，必要时重测。", "正确", ""),
    ("判断题", "修正结果可以替代实验条件判断；只要平台给出 K壁 和 KRe，就不需要关注 Re、壁效应和轨迹质量。", "错误", ""),
]


def set_font(paragraph, size=11, bold=False) -> None:
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    set_font(p, 16 if level == 1 else 13, True)


def add_para(doc: Document, text: str, size=11) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    set_font(p, size)


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia", "Microsoft YaHei")
    normal.font.size = Pt(11)


def build_lecture() -> Document:
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("基于 AI 视觉的落球法液体粘滞系数实验讲义")
    set_font(title, 18, True)
    add_para(doc, "本讲义用于平台准入学习和真实实验操作。学生应先阅读讲义并完成题目，再进入 AI 实验模块进行标定、背光测量、人工测量和报告生成。", 11)
    for heading, paragraphs in LECTURE_SECTIONS:
        add_heading(doc, heading, 1)
        for text in paragraphs:
            add_para(doc, text)
    return doc


def build_quiz() -> Document:
    doc = Document()
    style_doc(doc)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("落球法 AI 视觉实验准入题")
    set_font(title, 18, True)
    add_para(doc, "题型仅包含单项选择题和判断题。题目依据平台讲义设置，用于确认学生理解实验流程、标定要求、背光测量条件和结果复核方法。", 11)
    for index, (kind, title_text, answer, options) in enumerate(QUIZ_ITEMS, start=1):
        add_para(doc, f"{index}. [{kind}] {title_text}", 11)
        if options:
            for line in options.splitlines():
                add_para(doc, f"   {line}", 10.5)
        add_para(doc, f"   答案：{answer}", 10.5)
    return doc


def main() -> None:
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    build_lecture().save(LECTURE_PATH)
    build_quiz().save(QUIZ_PATH)
    print(LECTURE_PATH)
    print(QUIZ_PATH)


if __name__ == "__main__":
    main()
