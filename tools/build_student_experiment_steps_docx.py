from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("/Users/shi/Desktop/物创/落球法测液体粘滞系数实验_AI平台学生操作步骤.docx")

INK = RGBColor(20, 32, 48)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(90, 101, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
SOFT = "F7FAFC"
WHITE = "FFFFFF"


def set_font(run, size=11, color=INK, bold=False, name="Calibri", east_asia="Microsoft YaHei"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold


def para_format(p, before=0, after=6, line=1.25, align=None):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=90, start=130, bottom=90, end=130):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")
    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            width = widths[min(i, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_text(p, text, size=11, color=INK, bold=False):
    run = p.add_run(text)
    set_font(run, size=size, color=color, bold=bold)
    return run


def add_body(doc, text, after=6):
    p = doc.add_paragraph()
    para_format(p, after=after)
    add_text(p, text)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    para_format(p, before=18 if level == 1 else 12, after=8 if level == 1 else 5)
    for run in p.runs:
        set_font(run, size=16 if level == 1 else 13, color=BLUE if level == 1 else DARK_BLUE, bold=True)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        para_format(p, after=4)
        add_text(p, item)


def add_steps(doc, steps):
    for step in steps:
        p = doc.add_paragraph(style="List Number")
        para_format(p, after=5)
        add_text(p, step)


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table_geometry(table, [9360])
    cell = table.cell(0, 0)
    shade(cell, SOFT)
    p = cell.paragraphs[0]
    para_format(p, after=3)
    add_text(p, title, size=11, color=DARK_BLUE, bold=True)
    p2 = cell.add_paragraph()
    para_format(p2, after=0)
    add_text(p2, body, size=10.5)
    doc.add_paragraph()


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, header in enumerate(headers):
        table.cell(0, i).text = header
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    table_geometry(table, widths)
    for r_i, row in enumerate(table.rows):
        for cell in row.cells:
            shade(cell, LIGHT_BLUE if r_i == 0 else WHITE)
            for p in cell.paragraphs:
                para_format(p, after=0, line=1.15)
                for run in p.runs:
                    set_font(run, size=9.8 if len(headers) >= 4 else 10.3, color=INK, bold=r_i == 0)
    doc.add_paragraph()


def style_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def build_doc():
    doc = Document()
    style_doc(doc)

    header = doc.sections[0].header.paragraphs[0]
    para_format(header, after=0, line=1.0)
    add_text(header, "AI 视觉落球法实验平台 · 学生操作步骤", size=9.5, color=MUTED)
    footer = doc.sections[0].footer.paragraphs[0]
    para_format(footer, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(footer, "实验操作指导文件", size=9, color=MUTED)

    title = doc.add_paragraph()
    para_format(title, before=18, after=6, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(title, "落球法测液体粘滞系数实验", size=24, color=INK, bold=True)
    subtitle = doc.add_paragraph()
    para_format(subtitle, after=18, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_text(subtitle, "AI 平台学生操作步骤", size=15, color=MUTED, bold=True)

    add_callout(
        doc,
        "使用说明",
        "本文件只写学生实验操作流程。学生主流程为连接实时画面、现场释放小球、由平台自动追踪轨迹并输出结果；视频和 CSV 仅作为备用数据入口。",
    )

    add_heading(doc, "一、实验前准备", 1)
    add_bullets(
        doc,
        [
            "检查实验器材：量筒、待测液体、小球、释放装置、摄像设备或手机、补光灯、电脑。",
            "确认小球表面干净、无气泡附着，待测液体中无明显杂质。",
            "将量筒竖直固定，释放装置对准量筒中心位置，摄像设备保持稳定。",
            "如果使用手机拍摄，建议使用主摄，关闭明显的滤镜、美颜和过度防抖，保持画面清晰。",
            "准备需要填写的参数：液体名称、液体密度、小球半径、小球密度、量筒内径、液体深度、实验温度。",
        ]
    )

    add_heading(doc, "二、进入平台与完成预习", 1)
    add_steps(
        doc,
        [
            "打开平台首页，按首页提示进入学习流程。",
            "先阅读实验讲义，了解本次实验要测量的物理量和平台操作要求。",
            "阅读完成后进入试题测试页面，按题目要求完成测试。",
            "测试通过后进入实验大厅。",
            "在实验大厅中选择“AI 实验”模块，开始真实实验数据采集与分析。",
        ]
    )

    add_heading(doc, "三、填写实验参数", 1)
    add_body(doc, "进入 AI 实验页面后，先在左侧参数区填写本次实验的样本与仪器信息。填写完成后再开始采集数据。")
    add_table(
        doc,
        ["填写项", "填写内容", "注意事项"],
        [
            ["样本液体", "填写待测液体名称或样本编号。", "未知液体可以写“样本 1”“样本 A”等。"],
            ["液体密度", "填写待测液体密度。", "单位按页面要求填写，避免把 g/cm³ 与 kg/m³ 混用。"],
            ["小球半径", "填写小球半径。", "半径和直径不要填反。"],
            ["小球密度", "填写小球材料密度。", "钢球、玻璃球等材料密度应与实际材料一致。"],
            ["量筒内径", "填写量筒内部直径。", "不要填写外径。"],
            ["液体深度", "填写量筒中待测液体高度。", "读数时视线尽量与液面平齐。"],
            ["实验温度", "填写实验时液体温度。", "粘度对温度敏感，建议记录实际温度。"],
        ],
        [1900, 3600, 3860],
    )

    add_heading(doc, "四、选择数据采集方式", 1)
    add_body(doc, "本实验主流程采用实时追踪。学生应优先连接摄像头或手机实时画面，在平台中现场记录小球下落过程，并让平台自动生成轨迹、速度曲线和实验结果。视频文件和 CSV 文件只作为无法现场连接画面时的备用方式。")
    add_table(
        doc,
        ["方式", "定位", "学生要做什么"],
        [
            ["实时追踪", "正式实验主流程。", "连接实时画面，现场释放小球，由平台追踪小球位置并生成结果。"],
            ["实验视频", "备用入口。", "仅在无法现场连接实时画面时使用，选择已拍摄视频进行补充分析。"],
            ["轨迹 CSV", "备用入口。", "仅在已有外部轨迹数据时使用，确认字段和单位符合页面要求。"],
        ],
        [1500, 3700, 4160],
    )

    add_heading(doc, "五、实时追踪实验步骤", 1)
    add_steps(
        doc,
        [
            "在数据来源中选择“OpenCV 追踪”或实时追踪入口。",
            "点击“连接实时画面”，选择电脑摄像头、手机投屏画面或采集卡画面。",
            "调整量筒和摄像设备位置，使完整下落区域处于画面中，画面不要倾斜，背景尽量简洁。",
            "按平台页面提示完成实时画面设置；设置完成后不要移动摄像设备和量筒。",
            "点击“开始实时追踪”或“录制落球”，等待平台进入采集状态。",
            "轻轻释放小球，使小球沿量筒中心附近下落，避免用手推、甩动或让小球碰到筒壁。",
            "观察平台实时画面中小球轨迹是否被识别；若画面中小球丢失、贴壁或明显偏斜，应重新实验。",
            "小球完成一次下落后，点击“结束追踪”或按页面提示结束本次采集。",
            "等待平台生成位移曲线、速度曲线、终端速度和粘滞系数结果。",
        ]
    )

    add_heading(doc, "六、备用数据导入方式", 1)
    add_body(doc, "以下方式不是本实验的主流程。只有在实时追踪无法使用、需要补充分析或教师提供已有数据时，才使用视频或 CSV 导入。")
    add_heading(doc, "6.1 备用：导入实验视频", 2)
    add_steps(
        doc,
        [
            "选择“摄像机视频”数据来源。",
            "点击选择文件，上传本次实验拍摄的视频。",
            "等待平台读取视频时长、分辨率和文件信息。",
            "确认视频画面包含完整落球过程后，点击分析按钮。",
            "等待平台输出轨迹、曲线和结果。",
        ]
    )
    add_heading(doc, "6.2 备用：导入轨迹 CSV", 2)
    add_steps(
        doc,
        [
            "选择“轨迹 CSV”数据来源。",
            "上传包含时间和位置数据的 CSV 文件。",
            "确认 CSV 中至少包含 t 与 y 两列。",
            "点击“分析已选 CSV”。",
            "等待平台输出速度曲线与实验结果。",
        ]
    )

    add_heading(doc, "七、查看实验结果", 1)
    add_body(doc, "平台完成分析后，学生需要先检查结果是否完整，再决定是否记录或重做。")
    add_table(
        doc,
        ["检查项目", "应查看的内容", "处理方式"],
        [
            ["终端速度", "查看 vt 是否已输出，速度曲线是否出现稳定区段。", "若曲线波动过大，建议重新实验。"],
            ["粘滞系数", "查看 η 的数值和单位。", "记录平台输出结果。"],
            ["拟合质量", "查看 R²、异常点数量、追踪置信度等指标。", "若质量较差，检查视频清晰度和释放过程。"],
            ["实验条件", "查看 Re、壁效应风险等提示。", "若平台提示风险较高，在报告中说明或重做实验。"],
            ["不确定度", "按页面要求填写测量误差后查看最终表达。", "记录 η ± U 的结果。"],
        ],
        [1800, 4050, 3510],
    )

    add_heading(doc, "八、保存记录与提交", 1)
    add_steps(
        doc,
        [
            "确认本次结果可用后，保存或记录平台生成的实验数据。",
            "进入实验历史记录或数据检索区域，查看本次实验是否已保存。",
            "记录本次实验的关键参数、速度曲线、终端速度、粘滞系数和不确定度。",
            "如果进行了多次实验，选择质量较好的一组作为主要结果，其余数据可用于比较和误差分析。",
            "按教师要求提交实验记录、截图或平台导出的数据。",
        ]
    )

    add_heading(doc, "九、实验注意事项", 1)
    add_bullets(
        doc,
        [
            "实验过程中不要移动摄像设备、量筒和释放装置。",
            "释放小球时动作要轻，避免给小球明显初速度。",
            "小球若贴壁、偏斜严重或画面中途丢失，应重新实验。",
            "液体表面有气泡、反光强烈或背景杂乱时，会影响识别质量，应先调整环境。",
            "每次实验结束后，检查平台是否成功输出曲线和结果；未成功输出时不要直接提交。",
            "实验报告中应说明本次数据来源是实时追踪、视频还是 CSV；正式实验应优先使用实时追踪。",
        ]
    )

    add_callout(
        doc,
        "完成标准",
        "一次合格实验至少应包含：实时追踪得到的完整落球轨迹、速度曲线、终端速度、粘滞系数、质量指标、不确定度结果和实验记录。若其中任一项缺失，应根据平台提示重新采集。",
    )

    return doc


def main():
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
