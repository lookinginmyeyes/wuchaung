from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_PATH = Path("/Users/shi/Desktop/物创/基于AI视觉的落球法液体粘滞系数智能测量系统_平台讲义_送审稿.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(90, 101, 116)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], indent_dxa: int = 120) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri", east_asia="Microsoft YaHei") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_para_format(paragraph, before=0, after=6, line=1.25, align=None) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_run(paragraph, text, **kwargs):
    run = paragraph.add_run(text)
    set_font(run, **kwargs)
    return run


def add_body(doc: Document, text: str, after=6, before=0, align=None):
    p = doc.add_paragraph()
    set_para_format(p, before=before, after=after, align=align)
    add_run(p, text, size=11, color=INK)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 12), color=BLUE if level < 3 else DARK_BLUE, bold=True)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_para_format(p, after=4)
        add_run(p, item, size=11, color=INK)


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_para_format(p, after=4)
        add_run(p, item, size=11, color=INK)


def add_callout(doc: Document, title: str, body: str, fill: str = CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_para_format(p, after=3, line=1.25)
    add_run(p, title, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    set_para_format(p2, after=0, line=1.25)
    add_run(p2, body, size=10.5, color=INK)
    doc.add_paragraph()


def add_key_value_table(doc: Document, rows: list[tuple[str, str]], widths=(2700, 6660), header=None) -> None:
    table = doc.add_table(rows=1 if header else 0, cols=2)
    if header:
        table.cell(0, 0).text = header[0]
        table.cell(0, 1).text = header[1]
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
    set_table_geometry(table, list(widths), indent_dxa=120)
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if header and ri == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for p in cell.paragraphs:
                set_para_format(p, after=0, line=1.15)
                for run in p.runs:
                    set_font(run, size=10.3, color=INK, bold=(header and ri == 0) or ci == 0)
    doc.add_paragraph()


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths_dxa: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    set_table_geometry(table, widths_dxa, indent_dxa=120)
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            set_cell_shading(cell, LIGHT_BLUE if ri == 0 else WHITE)
            for p in cell.paragraphs:
                set_para_format(p, after=0, line=1.15)
                for run in p.runs:
                    set_font(run, size=9.7 if len(headers) >= 4 else 10.3, color=INK, bold=ri == 0)
    doc.add_paragraph()


def add_equation_block(doc: Document, title: str, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7FAFC")
    p = cell.paragraphs[0]
    set_para_format(p, after=4, line=1.15)
    add_run(p, title, size=10.5, bold=True, color=DARK_BLUE)
    for line in lines:
        eq = cell.add_paragraph()
        set_para_format(eq, after=2, line=1.15)
        add_run(eq, line, size=10.2, color=INK, name="Cambria Math", east_asia="Microsoft YaHei")
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    set_para_format(p, before=24, after=8, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p, "基于 AI 视觉的落球法液体粘滞系数智能测量系统", size=24, bold=True, color=INK)
    p2 = doc.add_paragraph()
    set_para_format(p2, after=20, line=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(p2, "平台讲义 · 送审稿", size=15, color=MUTED, bold=True)

    add_key_value_table(
        doc,
        [
            ("适用场景", "学生预习、平台准入测试前阅读、真实 AI 实验与虚拟仿真操作说明"),
            ("平台主线", "阅读讲义 → 完成试题 → 进入实验大厅 → AI 实验 / 虚拟仿真 / 误差诊断 / 历史记录"),
            ("核心目标", "用摄像机或轨迹数据获得小球终端速度，并结合修正 Stokes 公式计算待测液体粘滞系数"),
            ("审阅说明", "本文件为独立草稿，尚未替换平台内置讲义。确认后再导入平台。"),
        ],
        widths=(2100, 7260),
    )
    add_callout(
        doc,
        "讲义定位",
        "本讲义不把平台描述成一个“自动给答案”的黑箱，而是把物理模型、视觉标定、轨迹拟合、修正计算和误差判断串成一条可复核的实验流程。学生需要理解每一步为什么做、输入什么量、结果是否可信。",
    )


def build_doc() -> Document:
    doc = Document()
    style_document(doc)

    header = doc.sections[0].header.paragraphs[0]
    set_para_format(header, after=0, line=1.0)
    add_run(header, "落球法测液体粘滞系数实验 AI 学习平台 · 讲义送审稿", size=9.5, color=MUTED)
    footer = doc.sections[0].footer.paragraphs[0]
    set_para_format(footer, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_run(footer, "供平台讲义确认使用，未替换线上内容", size=9, color=MUTED)

    add_cover(doc)

    add_heading(doc, "一、学习目标与平台任务", 1)
    add_body(
        doc,
        "本平台服务于“落球法测量液体粘滞系数”的学习、实验和比赛展示。学生先通过讲义与准入测试理解核心物理量，再进入实验大厅使用 AI 实验、虚拟仿真、误差诊断和历史记录等模块完成完整实验闭环。",
    )
    add_bullets(
        doc,
        [
            "理解小球在液体中由加速到终端匀速的过程，知道终端速度 vt 是计算粘滞系数 η 的关键量。",
            "掌握 Stokes 公式的适用条件，并能解释低雷诺数、壁效应、释放初速度和匀速段选取对结果的影响。",
            "会使用平台的摄像机、视频或轨迹 CSV 入口，完成标定、追踪、速度拟合、粘度计算和不确定度分析。",
            "能用虚拟仿真先观察已知液体的小球速度形成过程，再回到真实实验中判断数据是否合理。",
        ]
    )
    add_matrix_table(
        doc,
        ["阶段", "学生任务", "平台输出"],
        [
            ["预习准入", "阅读讲义，完成试题；教师调试时可使用跳过测试入口。", "通过后进入实验大厅。"],
            ["AI 实验", "连接摄像头或导入视频/CSV，完成标定并记录落球过程。", "轨迹、速度曲线、vt、η、R²、Re、修正因子和诊断建议。"],
            ["虚拟仿真", "选择已知液体、容器和小球参数，生成速度曲线。", "终端速度形成过程、壁效应与雷诺数修正提示。"],
            ["复盘记录", "对照历史数据、不确定度和误差诊断。", "实验记录检索、偏差来源与重测建议。"],
        ],
        [1500, 4260, 3600],
    )

    add_heading(doc, "二、落球法的物理原理", 1)
    add_body(
        doc,
        "半径为 r 的光滑小球在静止液体中下落时，受到重力、浮力和液体粘滞阻力。开始阶段小球速度较小，阻力也较小，小球加速下落；随着速度增大，粘滞阻力增大。当重力等于浮力与阻力之和时，小球加速度趋近于零，进入终端匀速阶段。此时速度称为终端速度 vt。",
    )
    add_equation_block(
        doc,
        "理想 Stokes 公式",
        [
            "F = 6πηrv",
            "η0 = 2r²g(ρ球 − ρ液) / (9vt)",
            "其中 r 为小球半径，vt 为终端速度，ρ球 与 ρ液 分别为小球和液体密度，g = 9.80665 m/s²。",
        ],
    )
    add_callout(
        doc,
        "适用条件",
        "理想 Stokes 公式要求小球刚性、光滑、近似球形；液体为牛顿流体；下落过程处于低雷诺数、无明显涡流的黏性主导状态；小球已经进入终端匀速段；容器足够宽或已经进行壁效应修正；释放时尽量无初速度、无旋转、无明显偏斜。",
    )

    add_heading(doc, "三、壁效应与雷诺数修正", 1)
    add_body(
        doc,
        "真实实验使用量筒或圆筒容器，液体并不是无限广延介质。小球附近的液体流动会受到筒壁和液柱深度影响，因此平台在理想 Stokes 结果之外保留修正项，并用迭代方式求得最终粘滞系数。",
    )
    add_equation_block(
        doc,
        "平台采用的修正关系",
        [
            "Re = ρ液 · vt · (2r) / η",
            "K壁 = (1 + 2.4r/R)(1 + 3.3r/H)，其中 R 为量筒内半径，H 为待测液体深度。",
            "KRe = 1 + 3Re/16 − 19Re²/1080；二级项前为负号。",
            "η = η0 / (K壁 · KRe)，由于 Re 与 η 有关，平台会迭代更新 η 与 Re。",
        ],
    )
    add_key_value_table(
        doc,
        [
            ("D 与 R", "D 是量筒内径，R 是量筒内半径，R = D/2。平台界面若输入内径，会在计算中自动换算为半径。"),
            ("Re 判据", "教学实验通常要求 Re < 1；若希望更接近蠕动流条件，可取 Re < 0.1 作为更保守判据。"),
            ("壁效应判断", "球筒径比 2r/D 越小，壁面影响越弱；小球偏大或量筒偏窄时，需关注 K壁 和诊断提示。"),
            ("结果解释", "修正结果不是“把误差抹掉”，而是把已知系统偏差纳入模型，并提示实验条件是否仍可信。"),
        ],
        header=("概念", "说明"),
    )

    add_heading(doc, "四、AI 视觉测量的工作原理", 1)
    add_body(
        doc,
        "AI 实验模块把传统人工计时和目视选段改为视觉轨迹测量。摄像机固定拍摄小球下落过程，平台通过 OpenCV 识别小球中心位置，生成时间-位移轨迹，再从稳定匀速段拟合终端速度。",
    )
    add_numbers(
        doc,
        [
            "采集画面：可使用浏览器摄像头、手机投屏/采集卡、已拍摄视频，或直接导入轨迹 CSV。",
            "建立标定：通过中心标定棒多点点击，把像素坐标转换为真实纵向坐标。",
            "追踪小球：OpenCV 采用轮廓、圆检测或阈值分割得到小球中心，并输出 t、x、y、confidence 等轨迹字段。",
            "拟合速度：平台选择速度较稳定、异常点较少的区段，使用稳健加权线性拟合得到 vt。",
            "计算结果：代入修正 Stokes 公式得到 η，同时输出 Re、R²、匀速段离散度、追踪置信度和误差诊断。",
        ]
    )
    add_callout(
        doc,
        "关于 Huber 拟合",
        "平台所说的稳健拟合，是指在拟合匀速段时降低明显异常点的权重。这样反光、短暂遮挡或个别识别漂移不会直接把终端速度拉偏；但若异常点过多，系统仍会提示追踪质量不足，建议重测。",
    )

    add_heading(doc, "五、标定棒多点非线性修正", 1)
    add_body(
        doc,
        "真实拍摄会受到透视、镜头畸变、玻璃量筒折射和画面裁切影响。若要求学生填写折射率、玻璃壁厚、相机距离等参数，实际操作负担很大，也容易引入伪精度。因此平台采用更适合课堂和比赛展示的实测标定方法：用一根已知尺寸的标定棒直接建立像素坐标到真实坐标的映射。",
    )
    add_matrix_table(
        doc,
        ["标定对象", "推荐设置", "平台作用"],
        [
            ["相机标定板", "棋盘格或 ChArUco；同一手机、同一焦距、同一分辨率可复用。", "主要修正镜头畸变和广角边缘变形。"],
            ["中心标定棒", "直径 3 mm，长度 400 mm，建议 50 mm 一格，共点击 9 个点。", "建立 Yreal = f(ypixel) 的纵向映射。"],
            ["多点映射", "按 0、50、100 ... 400 mm 的真实顺序点击可见刻度。", "用分段插值吸收透视、折射和残余畸变造成的非等比例变化。"],
        ],
        [1900, 3560, 3900],
    )
    add_callout(
        doc,
        "为什么不是只用一个比例尺",
        "如果画面中上端、中部、下端的像素-长度比例并不相同，单一比例尺会把整段运动当成线性关系处理。多点标定相当于用实际刻度直接描出一条纵向坐标尺；相邻刻度间线性插值，但整体可以是非线性的，所以更适合量筒、液体和摄像机共同造成的复杂成像偏差。",
    )

    add_heading(doc, "六、平台操作步骤", 1)
    add_heading(doc, "6.1 预习准入", 2)
    add_numbers(
        doc,
        [
            "进入首页后，按页面提示向下拉动小球或点击进入按钮，进入讲义学习流程。",
            "阅读讲义，重点关注终端速度、Stokes 条件、壁效应、雷诺数和视觉标定。",
            "完成准入试题。正式教学中应通过测试后进入实验大厅；跳过测试按钮仅建议教师调试或展示时使用。",
        ]
    )
    add_heading(doc, "6.2 AI 实验：实时拍摄或数据导入", 2)
    add_numbers(
        doc,
        [
            "填写样本与仪器参数：液体名称或样本编号、液体密度、小球半径、小球密度、量筒内径、液体深度、温度等。",
            "选择数据来源：轨迹 CSV、摄像机视频、OpenCV 实时追踪或标定数据。",
            "若使用实时画面，先连接手机或摄像头；将标定棒置于量筒中心轴线上，点击“开始标定”，从上到下点击各刻度点。",
            "标定完成后移走标定棒，保持相机、量筒和释放装置位置不变；点击录制落球，释放小球，结束录制后平台提交 OpenCV 追踪。",
            "查看速度曲线与粘度结果，确认 vt、η、R²、Re、K壁、KRe、追踪置信度和诊断建议是否合理。",
            "在不确定度卡片中输入仪器分辨率或估读误差，生成 η ± U 的表达。",
        ]
    )
    add_heading(doc, "6.3 视频与 CSV 备选流程", 2)
    add_bullets(
        doc,
        [
            "视频导入适合已经用手机或相机拍完实验的情况。平台读取视频元信息后，后端取帧并追踪小球。",
            "CSV 导入适合已有轨迹点或外部软件导出的数据。必填列为 t、y，单位默认 t/s、y/m；可选列包括 x、confidence、measured_y、corrected_y。",
            "若 CSV 中已经包含 corrected_y，表示外部或平台标定后的位置；若只有 y，系统按输入比例或默认单位处理。",
        ]
    )
    add_heading(doc, "6.4 虚拟仿真与实验大厅", 2)
    add_bullets(
        doc,
        [
            "虚拟仿真用于已知液体：选择标准液体、容器条件和小球参数后，平台输出小球速度曲线与终端速度。",
            "初始扰动、介质阻尼扰动和速度采样稳定性属于实验质量与噪声控制量，用来观察释放质量、采样波动和识别稳定性对曲线的影响。",
            "实验大厅中的历史记录/数据检索用于回看每次实验参数、曲线、粘度结果和误差诊断，方便复测比较。",
        ]
    )

    add_heading(doc, "七、数据处理与不确定度", 1)
    add_body(
        doc,
        "平台首先从轨迹中寻找稳定匀速段，并对 y-t 曲线做稳健线性拟合。拟合斜率即终端速度 vt。随后代入修正 Stokes 公式得到 η，并在结果区给出可复核指标：R² 越接近 1，说明该段线性越好；Re 用于判断 Stokes 条件；匀速段离散度和追踪置信度用于判断视觉数据质量。",
    )
    add_equation_block(
        doc,
        "平台显示的不确定度传播式",
        [
            "Ur = Δη / η",
            "Ur = √{[(2/d − 2.4/((1 + 2.4d/D)D) − 1.6/((1 + 1.6d/H)H))Δd]²",
            "      + (Δt/t)² + (−Δl/l)²",
            "      + [2.4d/((1 + 2.4d/D)D²)ΔD]²",
            "      + [1.6d/((1 + 1.6d/H)H²)ΔH]²}",
        ],
    )
    add_key_value_table(
        doc,
        [
            ("d", "小球直径，平台由小球半径自动换算 d = 2r。"),
            ("D", "量筒内径，不是半径。若实测的是半径，应先换算为内径输入。"),
            ("H", "待测液体深度。"),
            ("l 与 t", "用于计算终端速度的有效下落距离和对应时间。"),
            ("Δd、ΔD、ΔH、Δl、Δt", "小球直径、量筒内径、液体深度、距离和时间的测量不确定度或估读误差。"),
            ("U(k=2)", "扩展不确定度，平台用 k=2 给出约 95% 置信水平下的结果表达。"),
        ],
        header=("符号", "含义"),
        widths=(1600, 7760),
    )

    add_heading(doc, "八、误差来源与复测建议", 1)
    add_matrix_table(
        doc,
        ["现象", "可能原因", "处理建议"],
        [
            ["速度曲线前段波动大", "释放带初速度、磁吸头或夹具干扰、小球旋转。", "只选终端稳定段；改进释放装置，保证小球从中心轴线附近静止释放。"],
            ["Re 偏高", "小球过大、速度过快、液体粘度偏低。", "换更小小球或更高粘度液体；报告中说明 Stokes 条件风险。"],
            ["壁效应明显", "量筒内径偏小或小球偏大，2r/D 偏高。", "使用更大内径量筒或更小球；保留壁效应修正并复核 K壁。"],
            ["追踪置信度低", "反光、气泡、遮挡、背景对比不足。", "调整补光和背景；减少液面反光；必要时使用视频重拍。"],
            ["标定后仍偏差大", "相机移动、标定棒未在量筒中心、点击顺序错误。", "重新固定相机；重新按从上到下的真实刻度点标定。"],
        ],
        [1700, 3500, 4160],
    )

    add_heading(doc, "九、实验记录表", 1)
    add_matrix_table(
        doc,
        ["项目", "记录内容", "填写示例"],
        [
            ["基本参数", "液体名称、温度、液体密度、小球半径、小球密度、量筒内径、液体深度。", "纯甘油 25℃；r=1.50 mm；D=35.0 mm；H=220 mm。"],
            ["标定记录", "相机是否固定、是否完成相机标定、标定棒长度、点击点数、平均比例。", "400 mm 标定棒；50 mm 间距；9 点。"],
            ["追踪结果", "vt、η、R²、Re、K壁、KRe、异常点数量、追踪置信度。", "以平台输出为准。"],
            ["不确定度", "Δd、ΔD、ΔH、Δl、Δt；最终 η ± U(k=2)。", "根据游标卡尺、刻度尺和计时/拟合数据填写。"],
            ["结论复盘", "本次结果是否满足低 Re 条件，是否需要重测，主要误差来源是什么。", "围绕 Re、壁效应、释放质量、标定质量说明。"],
        ],
        [1600, 5260, 2500],
    )

    add_heading(doc, "十、学生提交要求", 1)
    add_bullets(
        doc,
        [
            "提交平台导出的实验记录或截图，至少包含速度曲线、终端速度、粘滞系数、Re、壁效应修正因子和不确定度结果。",
            "说明数据来源：实时拍摄、视频导入或 CSV 导入；若使用实时拍摄，应说明标定棒点击点数和标定状态。",
            "用自己的语言解释本次实验是否满足 Stokes 公式的主要条件，若条件不理想，应写出改进方案。",
            "将真实实验结果与虚拟仿真或参考液体数据进行对照，讨论差异来自物理条件、视觉追踪还是人工测量误差。",
        ]
    )
    add_callout(
        doc,
        "最后提醒",
        "AI 视觉平台的价值不是替代物理判断，而是把原来难以看见的轨迹、速度、异常点和修正因子显示出来。学生最终仍需根据公式适用条件、实验操作和误差来源判断结果是否可靠。",
        fill=LIGHT_BLUE,
    )
    return doc


def main() -> None:
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc()
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    main()
